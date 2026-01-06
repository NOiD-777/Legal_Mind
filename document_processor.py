import os
import tempfile
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
import streamlit as st

class DocumentProcessor:
    """Handles extraction of text content from various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_txt_text
        }
    
    def extract_text(self, file_path: str, mime_type: str) -> Optional[str]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            mime_type: MIME type of the file
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            if mime_type in self.supported_formats:
                return self.supported_formats[mime_type](file_path)
            else:
                raise ValueError(f"Unsupported file format: {mime_type}")
        except Exception as e:
            st.error(f"Error extracting text from document: {str(e)}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for robustness"""
        text_content = ""
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n\n"
            
            # If pdfplumber didn't extract much text, try PyPDF2
            if len(text_content.strip()) < 100:
                text_content = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
            
            if not text_content.strip():
                raise ValueError("No text could be extracted from the PDF. The document may be image-based or encrypted.")
            
            return text_content.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in the Word document.")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'ascii', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            return content.strip()
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode the text file with any supported encoding.")
            
        except Exception as e:
            raise Exception(f"Failed to extract text from text file: {str(e)}")
    
    def get_document_stats(self, text: str) -> dict:
        """Get basic statistics about the document"""
        words = text.split()
        sentences = text.split('.')
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()])
        }
