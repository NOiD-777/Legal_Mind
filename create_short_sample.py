#!/usr/bin/env python3
"""
Create a shorter sample document that works within Gemini API rate limits
"""

from docx import Document

def create_short_contract():
    """Create a shorter contract for testing within API limits"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('SIMPLE SERVICE AGREEMENT', 0)
    title.alignment = 1
    
    doc.add_paragraph()
    
    # Much shorter contract with clear issues
    contract_text = """
SERVICE AGREEMENT

This Agreement is between TechCorp LLC ("Company") and Client Inc. ("Client").

1. SERVICES
Company will provide software development services. Company determines all work scope in its sole discretion. Client must accept all deliverables without warranty.

2. PAYMENT  
Client pays $100,000 upfront. No refunds under any circumstances. Company may add fees without notice.

3. INTELLECTUAL PROPERTY
All IP created becomes Company property exclusively, including Client's existing IP used in the project.

4. LIABILITY
Company has zero liability for any damages. Client indemnifies Company for all claims, including Company's own misconduct.

5. TERMINATION
Company may terminate anytime without cause or notice. Client cannot terminate under any circumstances.

6. DISPUTES
All disputes resolved in Delaware courts only. Client waives jury trial rights.

7. DATA
Company may use all Client data without restriction. Client waives all privacy rights under GDPR/CCPA.

Signed: _____________    Signed: _____________
Company                 Client
"""

    # Add paragraphs
    for paragraph in contract_text.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    filename = 'short_test_contract.docx'
    doc.save(filename)
    print(f"Short test contract created: {filename}")
    return filename

if __name__ == "__main__":
    create_short_contract()