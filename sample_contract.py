#!/usr/bin/env python3
"""
Script to create a sample legal document in DOCX format for testing
"""

from docx import Document
from docx.shared import Inches
import os

def create_sample_contract():
    """Create a sample service agreement with potential legal issues"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('SERVICE AGREEMENT', 0)
    title.alignment = 1  # Center alignment
    
    # Add some spacing
    doc.add_paragraph()
    
    # Contract content with intentional legal issues for testing
    contract_text = """
COMPREHENSIVE SOFTWARE DEVELOPMENT AND CONSULTING AGREEMENT

This Software Development and Consulting Agreement ("Agreement") is entered into on January 15, 2024, between TechCorp Solutions LLC, a Delaware limited liability company with principal offices at 123 Innovation Drive, Wilmington, DE 19801 ("Company" or "Contractor"), and ClientCorp Industries Inc., a California corporation with principal offices at 456 Business Boulevard, San Francisco, CA 94105 ("Client" or "Customer").

RECITALS

WHEREAS, Company is engaged in the business of providing custom software development, system integration, cybersecurity consulting, and digital transformation services to enterprise clients;

WHEREAS, Client operates a multi-location retail business with over 500 employees and requires comprehensive software solutions to modernize their inventory management, customer relationship management, and e-commerce platforms;

WHEREAS, Client has represented that it handles sensitive customer data including payment card information, personal identifiable information, and health-related data subject to various regulatory requirements;

WHEREAS, the parties desire to establish the terms and conditions under which Company will provide software development and consulting services to Client;

NOW, THEREFORE, in consideration of the mutual covenants, terms, and conditions contained herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the parties agree as follows:

1. SCOPE OF SERVICES AND DELIVERABLES

1.1 Services Overview
Company shall provide the following services ("Services"):
a) Custom software development for Client's inventory management system
b) Integration with existing legacy systems including SAP and Oracle databases
c) Development of customer-facing e-commerce platform with payment processing
d) Mobile application development for iOS and Android platforms
e) Cybersecurity assessment and implementation of security measures
f) Data migration services from legacy systems
g) Staff training and documentation
h) Ongoing maintenance and support services

1.2 Specific Deliverables
The specific scope of work will be determined by Company in its sole and absolute discretion based on Company's assessment of Client's needs. Client agrees to accept all deliverables as provided without warranty of any kind, express or implied. Company makes no representations regarding the suitability, functionality, or performance of any deliverables.

1.3 Timeline
All deliverables shall be completed according to Company's internal timeline. Company provides no guarantees regarding completion dates and shall not be liable for any delays regardless of cause.

2. COMPENSATION AND PAYMENT TERMS

2.1 Project Fee
Client shall pay Company a total project fee of $850,000 for all Services described herein. This amount is due in full within 10 days of execution of this Agreement, regardless of project completion status.

2.2 Additional Fees
Additional fees may be imposed by Company at its sole discretion without prior notice to Client, including but not limited to:
a) Scope changes (minimum $50,000 per change)
b) Technology updates during project ($25,000 per update)
c) Client-requested meetings ($5,000 per hour)
d) Documentation revisions ($10,000 per revision)
e) Emergency support ($15,000 per incident)

2.3 No Refunds
Under no circumstances will any portion of payments made to Company be refundable, regardless of project completion, deliverable quality, or termination circumstances.

2.4 Late Payment
Any payments not received within the specified timeframe shall accrue interest at 25% per annum, compounded monthly.

3. INTELLECTUAL PROPERTY RIGHTS

3.1 Ownership of Work Product
All intellectual property created during the term of this Agreement, including but not limited to software code, documentation, designs, processes, methodologies, and any derivative works thereof, shall become the exclusive property of Company. This includes any pre-existing intellectual property of Client that is incorporated into or modified during the project.

3.2 Client IP Assignment
Client hereby assigns, transfers, and conveys to Company all right, title, and interest in and to any intellectual property that Client owns or controls that relates to or is used in connection with the Services.

3.3 Waiver of Rights
Client expressly waives all rights to any work product, including moral rights, attribution rights, and any claims to authorship or inventorship.

4. CONFIDENTIALITY AND DATA HANDLING

4.1 Client Confidentiality Obligations
Client agrees to maintain in strict confidence all information disclosed by Company, including but not limited to business methods, pricing strategies, client lists, and technical information. This obligation shall remain in effect indefinitely, even after termination of this Agreement.

4.2 Company Rights
Company shall have no reciprocal confidentiality obligations whatsoever. Company may use, disclose, reproduce, and distribute any and all of Client's confidential information for any purpose, including commercial purposes, without restriction or compensation to Client.

4.3 Data Processing
Company may collect, process, store, and share any and all data provided by Client or generated during the provision of Services without limitation. Client hereby consents to unlimited data processing and expressly waives all privacy rights under applicable laws including but not limited to GDPR, CCPA, HIPAA, and any state privacy regulations.

5. LIABILITY, INDEMNIFICATION, AND RISK ALLOCATION

5.1 Complete Liability Limitation
Company shall have absolutely no liability for any damages whatsoever, whether direct, indirect, special, incidental, consequential, punitive, or exemplary damages, including but not limited to loss of profits, business interruption, loss of data, system failures, security breaches, or any other commercial damages or losses.

5.2 Broad Indemnification
Client agrees to indemnify, defend, and hold harmless Company, its officers, directors, employees, agents, and affiliates from and against any and all claims, demands, losses, costs, damages, liabilities, and expenses (including reasonable attorneys' fees) arising from or relating to:
a) Client's use of the Services or deliverables
b) Any breach of this Agreement by Client
c) Any negligent or wrongful acts or omissions by Company, including gross negligence and willful misconduct
d) Any third-party claims related to the Services
e) Any regulatory violations or compliance failures

5.3 Client Assumption of Risk
Client acknowledges and agrees that it is assuming all risks associated with the Services, including but not limited to risks of system failures, data loss, security breaches, regulatory violations, and business interruption.

6. TERMINATION PROVISIONS

6.1 Company Termination Rights
This Agreement may be terminated by Company at any time, with or without cause, and with or without notice to Client. Company may also suspend Services at any time for any reason.

6.2 Client Termination Restrictions
Client may not terminate this Agreement under any circumstances, including material breach by Company, failure to deliver Services, or fundamental frustration of purpose.

6.3 Effect of Termination
Upon termination for any reason:
a) All payments made to Company are immediately forfeited and non-refundable
b) Client must immediately cease all use of any deliverables
c) All intellectual property rights remain with Company
d) Client's indemnification obligations survive indefinitely

7. DISPUTE RESOLUTION AND GOVERNING LAW

7.1 Exclusive Jurisdiction
Any and all disputes, claims, or controversies arising out of or relating to this Agreement shall be resolved exclusively in the state and federal courts located in New Castle County, Delaware, regardless of where the dispute arises or where Client is located.

7.2 Waiver of Rights
Client hereby irrevocably waives any right to:
a) Trial by jury
b) Participate in class action lawsuits
c) Seek injunctive relief
d) Appeal any court decisions

7.3 Governing Law
This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware, without regard to conflict of law principles. Where Delaware law conflicts with any other applicable law, the interpretation most favorable to Company shall apply.

8. EMPLOYMENT AND LABOR PROVISIONS

8.1 Automatic Employment Transfer
Any employees of Client who work on or have knowledge of this project shall automatically become at-will employees of Company upon execution of this Agreement. Client waives all employment law protections for such employees and agrees that normal labor standards and employment laws do not apply.

8.2 Non-Solicitation
Client agrees not to solicit, hire, or engage any current or former employees or contractors of Company for a period of 10 years following termination of this Agreement.

9. REGULATORY COMPLIANCE AND WARRANTIES

9.1 Client Compliance Responsibility
Client is solely and exclusively responsible for ensuring compliance with all applicable laws, regulations, and industry standards in all jurisdictions where Client operates, including but not limited to:
a) Payment Card Industry Data Security Standard (PCI DSS)
b) Health Insurance Portability and Accountability Act (HIPAA)
c) California Consumer Privacy Act (CCPA)
d) General Data Protection Regulation (GDPR)
e) Sarbanes-Oxley Act (SOX)
f) All state and federal employment laws
g) All international trade and export control regulations

9.2 Disclaimer of Warranties
Company makes no representations, warranties, or guarantees regarding compliance with any laws, regulations, or industry standards. All Services are provided "AS IS" and "AS AVAILABLE" without any warranties whatsoever.

10. FORCE MAJEURE AND EXCUSED PERFORMANCE

10.1 Broad Force Majeure
Company's performance under this Agreement shall be excused for any reason including but not limited to:
a) Acts of God, natural disasters, or weather conditions
b) Government actions or regulations
c) Labor disputes or strikes
d) Supply chain disruptions
e) Company's business priorities or strategic decisions
f) Company's convenience or preference to work on other projects
g) Economic conditions or market changes

11. ASSIGNMENT AND MODIFICATION

11.1 Company Assignment Rights
Company may freely assign, transfer, or subcontract any or all of its rights and obligations under this Agreement without Client's consent.

11.2 Client Assignment Restrictions
Client may not assign, transfer, or delegate any of its rights or obligations under this Agreement without Company's prior written consent, which may be withheld in Company's sole discretion.

11.3 Unilateral Modification
Company may modify this Agreement at any time by providing written notice to Client. Continued performance by Client shall constitute acceptance of such modifications.

12. ADDITIONAL TERMS AND CONDITIONS

12.1 Entire Agreement
This Agreement constitutes the entire agreement between the parties and supersedes all prior negotiations, representations, or agreements relating to the subject matter hereof.

12.2 Severability Override
If any provision of this Agreement is deemed unenforceable, the entire Agreement shall become null and void, and Client shall remain liable for all payments made and obligations incurred.

12.3 Survival
All provisions that are beneficial to Company shall survive termination of this Agreement indefinitely, including but not limited to payment obligations, intellectual property assignments, indemnification provisions, and confidentiality obligations.

12.4 Attorney-Client Privilege
All communications between the parties are deemed to be confidential attorney-client privileged communications, regardless of the actual nature of such communications.

12.5 Binding Effect
This Agreement shall be binding upon Client's successors, assigns, heirs, personal representatives, and any entities that acquire Client or its assets.

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

TECHCORP SOLUTIONS LLC                    CLIENTCORP INDUSTRIES INC

By: _________________________________    By: _________________________________
Name: Jonathan P. Smith                   Name: [To be completed by Client]
Title: Chief Executive Officer            Title: [To be completed by Client]
Date: January 15, 2024                   Date: _______________________________

EXHIBIT A - TECHNICAL SPECIFICATIONS

1. SYSTEM REQUIREMENTS
- Integration with legacy SAP R/3 system (version unspecified)
- Oracle Database compatibility (version to be determined by Company)
- Support for 10,000+ concurrent users
- 99.99% uptime guarantee (Company makes no actual guarantees)
- Response time under 100ms for all operations
- Mobile responsive design for all devices
- Multi-language support (languages to be selected by Company)

2. SECURITY REQUIREMENTS
- Implementation of undefined "enterprise-grade" security
- Compliance with unspecified industry standards
- Data encryption using methods chosen by Company
- User authentication system (specifications not provided)
- Audit logging capabilities (scope undefined)

3. PERFORMANCE BENCHMARKS
- Page load times: Company will determine acceptable performance
- Database query optimization: As deemed necessary by Company
- Scalability testing: If Company deems appropriate
- Load balancing: Implementation at Company's discretion

EXHIBIT B - PAYMENT SCHEDULE ADDENDUM

Phase 1: $850,000 due upon signing (100% of total project cost)
All subsequent work performed at Company's discretion without additional Client input or approval required.

Note: This payment schedule is non-negotiable and non-refundable under any circumstances.
"""

    # Add the contract text
    for paragraph in contract_text.split('\n\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            if paragraph.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.')):
                p.style = 'List Number'
    
    # Save the document
    filename = 'sample_legal_contract.docx'
    doc.save(filename)
    print(f"Sample contract created: {filename}")
    return filename

def create_sample_employment_agreement():
    """Create a sample employment agreement with potential issues"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('EMPLOYMENT AGREEMENT', 0)
    title.alignment = 1
    
    doc.add_paragraph()
    
    contract_text = """
EMPLOYMENT AGREEMENT

This Employment Agreement ("Agreement") is made between MegaCorp Enterprises ("Company") and [Employee Name] ("Employee").

1. POSITION AND DUTIES
Employee shall serve as Software Developer and perform such duties as assigned by Company. Employee's duties may change at any time without notice or compensation adjustment.

2. COMPENSATION
Initial salary: $75,000 annually. Company may reduce salary at any time for any reason. No overtime compensation will be provided regardless of hours worked.

3. BENEFITS
Employee is not eligible for health insurance, retirement benefits, or paid time off during the first 24 months of employment.

4. NON-COMPETE
Employee agrees not to work for any technology company anywhere in the world for 5 years after termination, regardless of reason for termination.

5. INTELLECTUAL PROPERTY
All ideas, inventions, or creative works conceived by Employee during employment, including those created at home or on weekends, belong exclusively to Company.

6. TERMINATION
Employment is at-will. However, Employee must provide 90 days written notice to terminate, while Company may terminate immediately without cause or notice.

7. CONFIDENTIALITY
Employee must keep all Company information confidential for 10 years after termination, including publicly available information.

8. WORK SCHEDULE
Employee must be available 24/7 including weekends and holidays. Minimum 60 hours per week required.

9. EXPENSE REIMBURSEMENT
Employee is responsible for all work-related expenses including equipment, software, travel, and office supplies.

10. DISPUTE RESOLUTION
All disputes must be resolved through binding arbitration. Employee waives right to legal representation during arbitration proceedings.

Signed: _______________  Date: _______________
Employee

Signed: _______________  Date: _______________
Company Representative
"""

    for paragraph in contract_text.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    filename = 'sample_employment_agreement.docx'
    doc.save(filename)
    print(f"Sample employment agreement created: {filename}")
    return filename

if __name__ == "__main__":
    contract_file = create_sample_contract()
    employment_file = create_sample_employment_agreement()
    
    print(f"\nCreated sample documents:")
    print(f"1. {contract_file}")
    print(f"2. {employment_file}")
    print("\nThese documents contain intentional legal issues for testing the analyzer.")